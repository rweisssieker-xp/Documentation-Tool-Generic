"""
DOCX-Generator für Handbücher
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
from typing import List, Dict, Optional
from datetime import datetime
import os
from src.utils.logger import get_logger

logger = get_logger(__name__)


class DOCXBuilder:
    """Erstellt DOCX-Dokumente aus Handbuch-Daten"""
    
    def __init__(self, title: str = "Handbuch", author: str = None, version: str = "1.0", metadata: Optional[Dict] = None, template_config: Optional[Dict] = None):
        """
        Initialisiert den DOCX Builder
        
        Args:
            title: Titel des Dokuments
            author: Autor
            version: Version
            metadata: Erweiterte Metadaten (Abteilung, Projekt, Kontakt, etc.)
            template_config: Template-Konfiguration aus DocumentTemplate
        """
        self.document = Document()
        self.title = title
        self.author = author or os.getenv('USERNAME', 'Unbekannt')
        self.version = version
        self.metadata = metadata or {}
        self.template_config = template_config or {}
        
        self._setup_document()
        self._setup_styles()
    
    def _setup_document(self):
        """Konfiguriert Dokument-Grundlagen"""
        # Lade Formatierung aus Template falls verfügbar
        formatting = self.template_config.get('formatting', {})
        
        # Seitenränder
        sections = self.document.sections
        for section in sections:
            margin_top = formatting.get('margin_top', 1.0)
            margin_bottom = formatting.get('margin_bottom', 1.0)
            margin_left = formatting.get('margin_left', 1.0)
            margin_right = formatting.get('margin_right', 1.0)
            
            section.top_margin = Inches(margin_top)
            section.bottom_margin = Inches(margin_bottom)
            section.left_margin = Inches(margin_left)
            section.right_margin = Inches(margin_right)
    
    def _setup_styles(self):
        """Konfiguriert Dokumentstile basierend auf Template"""
        formatting = self.template_config.get('formatting', {})
        
        # Setze Schriftgrößen und -arten
        title_font_size = formatting.get('title_font_size', 24)
        heading_font_size = formatting.get('heading_font_size', 18)
        subheading_font_size = formatting.get('subheading_font_size', 14)
        body_font_size = formatting.get('body_font_size', 11)
        font_family = formatting.get('font_family', 'Calibri')
        heading_font_family = formatting.get('heading_font_family', 'Calibri')
        
        # Setze Farben
        color_scheme = formatting.get('color_scheme', {})
        primary_color = color_scheme.get('primary', '2E74B5')
        secondary_color = color_scheme.get('secondary', '5B9BD5')
        accent_color = color_scheme.get('accent', '70AD47')
        text_color = color_scheme.get('text', '000000')
        background_color = color_scheme.get('background', 'FFFFFF')
        
        # Aktualisiere Standardstile
        # Body Text Stil
        body_style = self.document.styles['Normal']
        body_style.font.name = font_family
        body_style.font.size = Pt(body_font_size)
        
        # Überschriften Stile
        if 'Heading 1' in self.document.styles:
            heading1_style = self.document.styles['Heading 1']
            heading1_style.font.name = heading_font_family
            heading1_style.font.size = Pt(heading_font_size)
            try:
                heading1_style.font.color.rgb = RGBColor.from_string(primary_color)
            except:
                heading1_style.font.color.rgb = RGBColor(0, 0, 0)
        
        if 'Heading 2' in self.document.styles:
            heading2_style = self.document.styles['Heading 2']
            heading2_style.font.name = heading_font_family
            heading2_style.font.size = Pt(subheading_font_size)
            try:
                heading2_style.font.color.rgb = RGBColor.from_string(secondary_color)
            except:
                heading2_style.font.color.rgb = RGBColor(0, 0, 0)
        
        # Zeilenabstand
        line_spacing = formatting.get('line_spacing', 1.15)
        paragraph_spacing_after = formatting.get('paragraph_spacing_after', 6)
        
        # Anwenden auf Standard-Paragraphen
        body_style.paragraph_format.line_spacing = line_spacing
        body_style.paragraph_format.space_after = Pt(paragraph_spacing_after)
    
    def add_title_page(self, title: str = None, subtitle: str = None):
        """
        Fügt Titelblatt hinzu
        
        Args:
            title: Titel
            subtitle: Untertitel
        """
        title = title or self.title
        
        # Hole Titelseiten-Einstellungen aus Template
        sections_config = self.template_config.get('sections', {}).get('title_page', {})
        formatting = self.template_config.get('formatting', {})
        title_font_size = formatting.get('title_font_size', 24)
        font_family = formatting.get('font_family', 'Calibri')
        
        # Titel
        title_paragraph = self.document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Pt(title_font_size)
        title_run.font.bold = True
        title_run.font.name = font_family
        
        self.document.add_paragraph()
        
        # Untertitel
        subtitle_text = subtitle or sections_config.get('subtitle', '')
        if subtitle_text and sections_config.get('subtitle', '').strip() != '':
            subtitle_paragraph = self.document.add_paragraph()
            subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_paragraph.add_run(subtitle_text)
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.name = font_family
        
        # Logo einfügen falls aktiviert
        if sections_config.get('include_logo', False) and sections_config.get('logo_path'):
            logo_path = sections_config.get('logo_path')
            if Path(logo_path).exists():
                try:
                    logo_paragraph = self.document.add_paragraph()
                    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logo_run = logo_paragraph.add_run()
                    logo_run.add_picture(str(logo_path), width=Inches(1.5))
                except Exception as e:
                    logger.warning(f"Fehler beim Einfügen des Logos: {e}", exc_info=True)
        
        # Metadaten
        self.document.add_paragraph()
        self.document.add_paragraph()
        
        meta_paragraph = self.document.add_paragraph()
        meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        meta_text = ""
        
        # Autor
        if sections_config.get('include_author', True):
            meta_text += f"Autor: {self.author}\n"
        
        # Datum
        if sections_config.get('include_date', True):
            meta_text += f"Erstellungsdatum: {datetime.now().strftime('%d.%m.%Y')}\n"
        
        meta_text += f"Version: {self.version}\n"
        
        # Erweiterte Metadaten aus Template-Metadaten
        template_metadata = self.template_config.get('metadata', {})
        if template_metadata.get('organization'):
            meta_text += f"Organisation: {template_metadata['organization']}\n"
        if template_metadata.get('department'):
            meta_text += f"Abteilung: {template_metadata['department']}\n"
        if template_metadata.get('project'):
            meta_text += f"Projekt: {template_metadata['project']}\n"
        if template_metadata.get('document_id'):
            meta_text += f"Dokument-ID: {template_metadata['document_id']}\n"
        
        # Erweiterte Metadaten aus Konstruktor
        if self.metadata.get('department') and not template_metadata.get('department'):
            meta_text += f"Abteilung: {self.metadata['department']}\n"
        if self.metadata.get('project') and not template_metadata.get('project'):
            meta_text += f"Projekt: {self.metadata['project']}\n"
        if self.metadata.get('contact'):
            meta_text += f"Kontakt: {self.metadata['contact']}\n"
        if self.metadata.get('document_id') and not template_metadata.get('document_id'):
            meta_text += f"Dokument-ID: {self.metadata['document_id']}\n"
        
        meta_run = meta_paragraph.add_run(meta_text)
        meta_run.font.size = Pt(10)
        meta_run.font.name = font_family
        
        # Seitenumbruch
        self.document.add_page_break()
    
    def add_table_of_contents(self):
        """Fügt Inhaltsverzeichnis hinzu"""
        toc_paragraph = self.document.add_paragraph()
        toc_run = toc_paragraph.add_run("Inhaltsverzeichnis")
        toc_run.font.size = Pt(18)
        toc_run.font.bold = True
        
        self.document.add_paragraph()
        
        # Erstelle manuelles Inhaltsverzeichnis
        # Word generiert automatisch ein TOC-Feld, aber python-docx unterstützt das nicht direkt
        # Wir erstellen daher ein manuelles TOC basierend auf den Überschriften
        
        # Sammle alle Überschriften während des Dokumentaufbaus
        # Diese Methode wird nach dem Hinzufügen aller Inhalte aufgerufen
        # Für jetzt: Platzhalter, wird später durch echte TOC-Generierung ersetzt
        self.document.add_paragraph("Inhaltsverzeichnis wird automatisch generiert...")
        self.document.add_page_break()
    
    def update_table_of_contents(self):
        """
        Aktualisiert das Inhaltsverzeichnis mit allen Überschriften
        Diese Methode sollte nach dem Hinzufügen aller Inhalte aufgerufen werden
        """
        # Suche nach dem TOC-Platzhalter
        # Erstelle neues TOC basierend auf Überschriften-Level
        headings = []
        for paragraph in self.document.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                level = 1
                if paragraph.style.name == 'Heading 1':
                    level = 1
                elif paragraph.style.name == 'Heading 2':
                    level = 2
                elif paragraph.style.name == 'Heading 3':
                    level = 3
                
                text = paragraph.text.strip()
                if text:
                    headings.append({'level': level, 'text': text})
        
        # Erstelle TOC-Einträge
        # Finde TOC-Paragraph und ersetze ihn
        for i, paragraph in enumerate(self.document.paragraphs):
            if "Inhaltsverzeichnis wird automatisch generiert" in paragraph.text:
                # Ersetze durch echtes TOC
                paragraph.clear()
                for heading in headings:
                    indent = "  " * (heading['level'] - 1)
                    run = paragraph.add_run(f"{indent}{heading['text']}\n")
                    run.font.size = Pt(11)
                break
    
    def add_introduction(self, text: str):
        """
        Fügt Einleitung hinzu
        
        Args:
            text: Einleitungstext
        """
        sections_config = self.template_config.get('sections', {})
        intro_config = sections_config.get('introduction', {})
        intro_title = intro_config.get('title', 'Einleitung')
        
        # Prüfe, ob Einleitung eingeschlossen werden soll
        structure = self.template_config.get('structure', {})
        if not structure.get('include_introduction', True):
            return
        
        heading = self.document.add_heading(intro_title, level=1)
        
        paragraph = self.document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def add_step(self, step: Dict, include_screenshot: bool = True):
        """
        Fügt einen Schritt hinzu
        
        Args:
            step: Schritt-Dictionary
            include_screenshot: Ob Screenshot eingefügt werden soll
        """
        step_number = step.get('step_number', 0)
        window_title = step.get('window_title', 'Unbekannt')
        description = step.get('description', '')
        timestamp = step.get('timestamp', 'N/A')
        
        # Hole Schritt-Einstellungen aus Template
        sections_config = self.template_config.get('sections', {})
        steps_config = sections_config.get('steps', {})
        include_step_numbering = steps_config.get('include_step_numbering', True)
        include_window_title = steps_config.get('include_window_title', True)
        include_timestamp = steps_config.get('include_timestamp', False)
        include_screenshot_caption = steps_config.get('include_screenshot_caption', True)
        screenshot_alignment = steps_config.get('screenshot_alignment', 'center')
        screenshot_size = steps_config.get('screenshot_size', 'medium')  # small, medium, large, original
        
        # Überschrift
        heading_parts = []
        if include_step_numbering:
            heading_parts.append(f"Schritt {step_number}")
        if include_window_title:
            heading_parts.append(window_title)
        
        heading_text = ": ".join(heading_parts) if heading_parts else "Schritt"
        heading = self.document.add_heading(heading_text, level=2)
        
        # Screenshot
        if include_screenshot:
            screenshot_path = Path(step.get('screenshot_path', ''))
            if screenshot_path.exists():
                try:
                    paragraph = self.document.add_paragraph()
                    
                    # Setze Alignment
                    if screenshot_alignment == 'center':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif screenshot_alignment == 'left':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif screenshot_alignment == 'right':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    run = paragraph.add_run()
                    
                    # Setze Bildgröße basierend auf Template
                    width_mapping = {
                        'small': Inches(4),
                        'medium': Inches(6),
                        'large': Inches(8),
                        'original': None  # Keine Größenänderung
                    }
                    
                    if screenshot_size in width_mapping and width_mapping[screenshot_size]:
                        run.add_picture(str(screenshot_path), width=width_mapping[screenshot_size])
                    else:
                        # Bei 'original' oder unbekannten Werten, füge Bild in Originalgröße ein
                        run.add_picture(str(screenshot_path))
                    
                    # Bildunterschrift
                    if include_screenshot_caption:
                        caption_paragraph = self.document.add_paragraph()
                        if screenshot_alignment == 'center':
                            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif screenshot_alignment == 'left':
                            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif screenshot_alignment == 'right':
                            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            
                        caption_run = caption_paragraph.add_run(f"Abbildung {step_number}: {window_title}")
                        caption_run.font.size = Pt(9)
                        caption_run.italic = True
                    
                    self.document.add_paragraph()
                
                except Exception as e:
                    logger.warning(f"Fehler beim Einfügen des Screenshots: {e}", exc_info=True)
        
        # Beschreibung
        if description:
            paragraph = self.document.add_paragraph(description)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Metadaten (optional, klein)
        if include_timestamp:
            meta_text = f"Zeitstempel: {timestamp}"
            meta_paragraph = self.document.add_paragraph()
            meta_run = meta_paragraph.add_run(meta_text)
            meta_run.font.size = Pt(8)
            meta_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def add_steps(self, steps: List[Dict], include_screenshots: bool = True):
        """
        Fügt mehrere Schritte hinzu
        
        Args:
            steps: Liste von Schritt-Dictionaries
            include_screenshots: Ob Screenshots eingefügt werden sollen
        """
        for step in steps:
            self.add_step(step, include_screenshots)
    
    def add_conclusion(self, text: str):
        """
        Fügt Fazit hinzu
        
        Args:
            text: Fazit-Text
        """
        sections_config = self.template_config.get('sections', {})
        conclusion_config = sections_config.get('conclusion', {})
        conclusion_title = conclusion_config.get('title', 'Fazit')
        
        # Prüfe, ob Fazit eingeschlossen werden soll
        structure = self.template_config.get('structure', {})
        if not structure.get('include_conclusion', True):
            return
        
        heading = self.document.add_heading(conclusion_title, level=1)
        
        paragraph = self.document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def add_security_notes(self, text: str):
        """
        Fügt Sicherheitshinweise hinzu
        
        Args:
            text: Sicherheitshinweise-Text
        """
        sections_config = self.template_config.get('sections', {})
        security_config = sections_config.get('security_notes', {})
        security_title = security_config.get('title', 'Sicherheitshinweise')
        
        # Prüfe, ob Sicherheitshinweise eingeschlossen werden sollen
        structure = self.template_config.get('structure', {})
        if not structure.get('include_security_notes', True):
            return
        
        heading = self.document.add_heading(security_title, level=1)
        
        paragraph = self.document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def add_troubleshooting(self, items: List[Dict[str, str]]):
        """
        Fügt Troubleshooting-Sektion hinzu
        
        Args:
            items: Liste von Problem-Lösung-Paaren
        """
        sections_config = self.template_config.get('sections', {})
        troubleshooting_config = sections_config.get('troubleshooting', {})
        troubleshooting_title = troubleshooting_config.get('title', 'Fehlerbehebung')
        
        # Prüfe, ob Troubleshooting eingeschlossen werden soll
        structure = self.template_config.get('structure', {})
        if not structure.get('include_troubleshooting', False):
            return
        
        heading = self.document.add_heading(troubleshooting_title, level=1)
        
        for item in items:
            problem = item.get('problem', '')
            solution = item.get('solution', '')
            
            problem_paragraph = self.document.add_paragraph(f"Problem: {problem}", style='List Bullet')
            solution_paragraph = self.document.add_paragraph(f"Lösung: {solution}")
            solution_paragraph.style = 'List Bullet 2'
    
    def save(self, output_path: Path) -> Path:
        """
        Speichert das Dokument
        
        Args:
            output_path: Ausgabepfad
            
        Returns:
            Pfad zur gespeicherten Datei
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Stelle sicher, dass Extension .docx ist
        if output_path.suffix != '.docx':
            output_path = output_path.with_suffix('.docx')
        
        # Aktualisiere Inhaltsverzeichnis vor dem Speichern
        self.update_table_of_contents()
        
        self.document.save(str(output_path))
        return output_path


